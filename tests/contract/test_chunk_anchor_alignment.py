from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from doc_chunk.api import chunk_document, extract_file, extract_outline

FIXTURES = Path(__file__).resolve().parents[1] / "fixtures"
NO_HEADING_DOCX = FIXTURES / "no_heading_style.docx"


@pytest.fixture(scope="module", autouse=True)
def ensure_no_heading_fixture() -> None:
    if NO_HEADING_DOCX.exists():
        return
    FIXTURES.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_paragraph("1. 投标须知")
    doc.add_paragraph("须知正文第一段。")
    doc.add_paragraph("2. 技术方案")
    doc.add_paragraph("方案描述。")
    doc.add_paragraph("2.1 系统架构")
    doc.add_paragraph("架构说明。")
    doc.add_paragraph("3. 报价说明")
    doc.add_paragraph("报价表格见下。")
    doc.save(NO_HEADING_DOCX)


def test_no_heading_docx_chunk_count_matches_outline(tmp_path: Path) -> None:
    ws = tmp_path / "no-heading-ws"
    extract_file(NO_HEADING_DOCX, ws, overwrite=True)
    outline = extract_outline(ws)
    chunk_document(ws, use_refined=False)
    import json

    index = json.loads((ws / "chunks" / "index.json").read_text(encoding="utf-8"))
    outline_nodes = list(outline.nodes)
    main_chunks = [c for c in index["chunks"] if c["title"] != "Preface"]
    ratio = len(main_chunks) / max(len(outline_nodes), 1)
    assert ratio >= 0.8, f"chunks={len(main_chunks)} outline={len(outline_nodes)}"
    assert len(main_chunks) > 1
    for entry in main_chunks:
        assert entry.get("original_node_ids"), entry["title"]
