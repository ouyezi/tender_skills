from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    path = tmp_path / "bid.docx"
    doc = Document()
    doc.add_heading("投标人须知", level=1)
    doc.add_paragraph("废标条款示例：未按规定递交投标文件。")
    doc.save(path)
    return path


@pytest.fixture
def sample_workspace(tmp_path: Path, sample_docx: Path) -> Path:
    from doc_chunk.api import run_pipeline

    ws = tmp_path / "ws"
    run_pipeline(sample_docx, ws, overwrite=True, skip_refine=True, skip_enrich=True)
    return ws
