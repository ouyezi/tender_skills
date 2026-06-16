"""Generate outline anchor fixtures for contract tests."""

from __future__ import annotations

import base64
import json
from pathlib import Path

from docx import Document

FIXTURES = Path(__file__).resolve().parents[1] / "fixtures"
MINIMAL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg=="
)


def ensure_anchor_fixtures() -> None:
    FIXTURES.mkdir(parents=True, exist_ok=True)
    png_path = FIXTURES / "minimal.png"
    if not png_path.exists():
        png_path.write_bytes(MINIMAL_PNG)

    image_docx = FIXTURES / "outline_anchor_on_image.docx"
    if not image_docx.exists():
        doc = Document()
        doc.add_picture(str(png_path))
        doc.add_paragraph("封面")
        doc.add_paragraph("第二章")
        doc.add_paragraph("第二章正文。")
        doc.save(image_docx)

    table_docx = FIXTURES / "outline_anchor_on_table.docx"
    if not table_docx.exists():
        doc = Document()
        table = doc.add_table(rows=2, cols=1)
        table.cell(0, 0).text = "封面"
        table.cell(1, 0).text = "表格说明"
        doc.add_paragraph("第二章")
        doc.add_paragraph("第二章正文。")
        doc.save(table_docx)


def patch_outline_anchor_on_block(ws: Path, *, outline_titles: list[str], block_index: int) -> None:
    outline_path = ws / "outline.json"
    outline = json.loads(outline_path.read_text(encoding="utf-8"))
    nodes = outline.get("nodes", [])
    if not nodes:
        nodes = [
            {
                "node_id": "n1",
                "title": outline_titles[0],
                "level": 1,
                "parent_id": None,
                "sort_order": 0,
                "anchor": {"char_start": 0, "block_start": block_index},
            },
            {
                "node_id": "n2",
                "title": outline_titles[1] if len(outline_titles) > 1 else "第二章",
                "level": 1,
                "parent_id": None,
                "sort_order": 1,
                "anchor": {"char_start": 0, "block_start": block_index + 1},
            },
        ]
    else:
        nodes[0]["anchor"] = {"char_start": 0, "block_start": block_index, "block_index": block_index}
        if len(nodes) < 2:
            nodes.append(
                {
                    "node_id": "n2",
                    "title": "第二章",
                    "level": 1,
                    "parent_id": None,
                    "sort_order": 1,
                    "anchor": {"char_start": 0, "block_start": block_index + 2},
                }
            )
    outline["nodes"] = nodes
    outline["strategy"] = "toc"
    outline_path.write_text(json.dumps(outline, ensure_ascii=False, indent=2), encoding="utf-8")
