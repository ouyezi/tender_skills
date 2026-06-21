from __future__ import annotations

import base64
import copy
from pathlib import Path

import pytest
from docx import Document


_TINY_PNG_BASE64 = (
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/w8AAgMBgN3k5+EAAAAASUVORK5CYII="
)


def _build_sample_docx(target: Path, *, include_image: bool) -> Path:
    doc = Document()
    doc.add_heading("示例标题", level=1)
    doc.add_paragraph("这是一段用于测试的正文。")
    if include_image:
        image_path = target.parent / "inline.png"
        image_path.write_bytes(base64.b64decode(_TINY_PNG_BASE64))
        doc.add_picture(str(image_path))
    doc.save(target)
    return target


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    return _build_sample_docx(tmp_path / "sample.docx", include_image=False)


@pytest.fixture
def sample_docx_with_image(tmp_path: Path) -> Path:
    return _build_sample_docx(tmp_path / "sample_with_image.docx", include_image=True)


@pytest.fixture
def sample_docx_with_duplicate_embed(tmp_path: Path) -> Path:
    path = tmp_path / "sample_duplicate_embed.docx"
    image_path = tmp_path / "inline.png"
    image_path.write_bytes(base64.b64decode(_TINY_PNG_BASE64))
    doc = Document()
    doc.add_paragraph("授权书附件")
    doc.add_picture(str(image_path))
    duplicate_paragraph = copy.deepcopy(doc.paragraphs[-1]._element)
    doc.element.body.append(duplicate_paragraph)
    doc.save(path)
    return path


@pytest.fixture
def sample_docx_with_duplicate_blip_in_paragraph(tmp_path: Path) -> Path:
    path = tmp_path / "sample_duplicate_blip_in_paragraph.docx"
    image_path = tmp_path / "inline.png"
    image_path.write_bytes(base64.b64decode(_TINY_PNG_BASE64))
    doc = Document()
    doc.add_picture(str(image_path))
    blips = [node for node in doc.paragraphs[-1]._element.iter() if node.tag.endswith("}blip")]
    if blips:
        blips[0].getparent().append(copy.deepcopy(blips[0]))
    doc.save(path)
    return path


@pytest.fixture
def sample_docx_with_reused_image_across_paragraphs(tmp_path: Path) -> Path:
    path = tmp_path / "sample_reused_image_across_paragraphs.docx"
    image_path = tmp_path / "inline.png"
    image_path.write_bytes(base64.b64decode(_TINY_PNG_BASE64))
    doc = Document()
    doc.add_paragraph("授权书")
    doc.add_picture(str(image_path))
    doc.add_paragraph("身份证明")
    doc.add_picture(str(image_path))
    doc.save(path)
    return path


@pytest.fixture
def sample_docx_with_reused_image_in_body_and_table(tmp_path: Path) -> Path:
    path = tmp_path / "sample_reused_image_in_body_and_table.docx"
    image_path = tmp_path / "inline.png"
    image_path.write_bytes(base64.b64decode(_TINY_PNG_BASE64))
    doc = Document()
    doc.add_paragraph("授权书正文")
    doc.add_picture(str(image_path))
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].add_run().add_picture(str(image_path))
    doc.save(path)
    return path


@pytest.fixture
def merged_colspan_docx(tmp_path: Path) -> Path:
    path = tmp_path / "merged_colspan.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(0, 0).text = "姓名"
    table.cell(0, 2).text = "角色"
    table.cell(1, 0).merge(table.cell(1, 1))
    table.cell(1, 0).text = "刘敏"
    table.cell(1, 2).text = "开发"
    doc.save(path)
    return path


@pytest.fixture
def personnel_dual_row_docx(tmp_path: Path) -> Path:
    path = tmp_path / "personnel_dual_row.docx"
    doc = Document()
    table = doc.add_table(rows=4, cols=4)
    headers1 = ["姓名", "本项目工作角色", "性别", "学历"]
    values1 = ["刘敏", "开发工程师", "男", "本科"]
    headers2 = ["级别", "年龄", "毕业学校", "从业年限"]
    values2 = ["高级Java工程师", "35", "承德石油学院", "9+"]
    for c, h in enumerate(headers1):
        table.cell(0, c).text = h
    for c, v in enumerate(values1):
        table.cell(1, c).text = v
    for c, h in enumerate(headers2):
        table.cell(2, c).text = h
    for c, v in enumerate(values2):
        table.cell(3, c).text = v
    doc.save(path)
    return path
