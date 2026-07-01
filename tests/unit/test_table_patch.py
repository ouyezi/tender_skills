from __future__ import annotations

from pathlib import Path

from docx import Document

from doc_chunk.extract.block_index import BlockAccumulator, write_accumulator_markdown, write_content_blocks
from doc_chunk.models.table_model import TableSidecar
from doc_chunk.table.patch import patch_docx_tables
from doc_chunk.workspace.layout import OutputWorkspace


def _setup_workspace_with_sidecar(tmp_path: Path) -> OutputWorkspace:
    ws = OutputWorkspace.create(tmp_path / "ws", overwrite=True)
    ref = "tables/t0000.json"
    sidecar = TableSidecar(
        block_index=0,
        layout_type="simple",
        grid_width=2,
        grid={
            "rows": [
                {
                    "cells": [
                        {"text": "H1", "colspan": 1, "rowspan": 1},
                        {"text": "H2", "colspan": 1, "rowspan": 1},
                    ]
                },
                {
                    "cells": [
                        {"text": "V1", "colspan": 1, "rowspan": 1},
                        {"text": "V2", "colspan": 1, "rowspan": 1},
                    ]
                },
            ]
        },
        logical_rows=[["H1", "H2"], ["V1", "V2"]],
        markdown="| H1 | H2 |\n| --- | --- |\n| V1 | V2 |",
        llm_text="table",
    )
    (ws.root / ref).write_text(sidecar.model_dump_json(indent=2), encoding="utf-8")
    acc = BlockAccumulator()
    acc.add_table(sidecar.markdown, table_ref=ref)
    write_accumulator_markdown(ws, acc)
    write_content_blocks(ws, acc.finalize())
    return ws


def test_patch_docx_tables_replaces_placeholder_with_word_table(tmp_path: Path) -> None:
    ws = _setup_workspace_with_sidecar(tmp_path)
    md = ws.content_path.read_text(encoding="utf-8")

    doc = Document()
    for line in md.splitlines():
        doc.add_paragraph(line)
    assert len(doc.tables) == 0

    result = patch_docx_tables(doc, ws)
    assert result.patched_count == 1
    assert len(doc.tables) == 1
    assert doc.tables[0].cell(0, 0).text == "H1"
    assert "<!-- table-ref:" not in "\n".join(p.text for p in doc.paragraphs)
    assert not any(p.text.strip().startswith("|") for p in doc.paragraphs)
