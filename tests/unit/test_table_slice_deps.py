from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from doc_chunk.table.slice_deps import (
    collect_embed_relationship_ids,
    collect_style_ids_from_tbl,
    iter_tbl_elements,
)


@pytest.fixture
def docx_with_table_and_image(tmp_path: Path) -> Path:
    path = tmp_path / "tbl_img.docx"
    img = tmp_path / "cell.png"
    img.write_bytes(
        b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01"
        b"\x00\x00\x00\x01\x08\x06\x00\x00\x00\x1f\x15\xc4\x89"
        b"\x00\x00\x00\nIDATx\x9cc\x00\x01\x00\x00\x05\x00\x01"
        b"\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82"
    )
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].add_run().add_picture(str(img))
    doc.save(path)
    return path


def test_collect_embed_relationship_ids_finds_cell_image(docx_with_table_and_image: Path) -> None:
    doc = Document(docx_with_table_and_image)
    tbl = doc.tables[0]._tbl
    rids = collect_embed_relationship_ids(tbl)
    assert len(rids) >= 1


def test_collect_style_ids_from_tbl_returns_set(docx_with_table_and_image: Path) -> None:
    doc = Document(docx_with_table_and_image)
    tbl = doc.tables[0]._tbl
    styles = collect_style_ids_from_tbl(tbl)
    assert isinstance(styles, set)


def test_iter_tbl_elements_yields_tbl_only(docx_with_table_and_image: Path) -> None:
    doc = Document(docx_with_table_and_image)
    tags = [el.tag.split("}")[-1] for el in iter_tbl_elements(doc.tables[0]._tbl)]
    assert "tbl" in tags
    assert "tr" in tags
    assert "tc" in tags
