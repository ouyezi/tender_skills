from __future__ import annotations

from pathlib import Path

from docx import Document

from doc_chunk.api import extract_file, extract_outline
from doc_chunk.extract.promote_headings import parse_content_heading_line
from doc_chunk.models.outline import OutlineTree


def test_parse_content_heading_line_cn_enum() -> None:
    assert parse_content_heading_line("二、百福得服务方案介绍") == (1, "二、百福得服务方案介绍")
    assert parse_content_heading_line("  三、项目服务保障能力  ") == (1, "三、项目服务保障能力")
    assert parse_content_heading_line("一、员工投诉反馈机制：") is None
    assert parse_content_heading_line("1、景区门票：接入全国8000多家") is None
    assert parse_content_heading_line("景区门票：接入全国8000多家旅游景点") is None
    assert parse_content_heading_line("1.口腔健康") is None
    assert parse_content_heading_line("1. 技术方案") == (1, "技术方案")
    assert parse_content_heading_line("1\u3000先进的订单管理系统：很长") is None
    assert parse_content_heading_line("1.1企业介绍") == (2, "企业介绍")
    assert parse_content_heading_line("普通段落。") is None


def test_promote_headings_auto_promotes_cn_enum(tmp_path: Path) -> None:
    docx_path = tmp_path / "cn_enum.docx"
    doc = Document()
    doc.add_heading("一、企业简介及资质", level=1)
    doc.add_paragraph("1.1 企业介绍")
    doc.add_paragraph("二、百福得服务方案介绍")
    doc.add_paragraph("2.1 企业福利管理的痛点及挑战")
    doc.save(docx_path)

    workspace = tmp_path / "ws"
    extract_file(docx_path, workspace, overwrite=True, promote_headings="auto")
    content_md = (workspace / "content.md").read_text(encoding="utf-8")
    assert "# 一、企业简介及资质" in content_md
    assert "# 二、百福得服务方案介绍" in content_md

    extract_outline(workspace)
    outline = OutlineTree.model_validate_json((workspace / "outline.json").read_text(encoding="utf-8"))
    roots = [node for node in outline.nodes if node.parent_id is None]
    assert len(roots) == 2
    root_titles = {node.title for node in roots}
    assert "一、企业简介及资质" in root_titles
    assert "二、百福得服务方案介绍" in root_titles

    by_title = {node.title: node for node in outline.nodes}
    assert by_title["企业福利管理的痛点及挑战"].parent_id == by_title["二、百福得服务方案介绍"].node_id
    assert by_title["企业介绍"].parent_id == by_title["一、企业简介及资质"].node_id


def test_promote_headings_keeps_local_cn_enum_series_as_paragraphs(tmp_path: Path) -> None:
    docx_path = tmp_path / "local_enum_series.docx"
    doc = Document()
    doc.add_paragraph("2.2.5.2 员工保险")
    doc.add_paragraph("一、员工投诉反馈机制：")
    doc.add_paragraph("投诉处理说明。")
    doc.add_paragraph("二、医疗险索赔流程：")
    doc.add_paragraph("索赔流程说明。")
    doc.add_paragraph("三、索赔及服务问答")
    doc.add_paragraph("问答内容。")
    doc.add_paragraph("四、应急措施")
    doc.add_paragraph("应急说明。")
    doc.save(docx_path)

    workspace = tmp_path / "ws"
    extract_file(docx_path, workspace, overwrite=True, promote_headings="auto")
    content_md = (workspace / "content.md").read_text(encoding="utf-8")
    assert "三、索赔及服务问答" in content_md
    assert "# 三、索赔及服务问答" not in content_md
    assert "# 四、应急措施" not in content_md
    assert "#### 员工保险" in content_md or "## 员工保险" in content_md

    extract_outline(workspace)
    outline = OutlineTree.model_validate_json((workspace / "outline.json").read_text(encoding="utf-8"))
    roots = [node for node in outline.nodes if node.parent_id is None]
    root_titles = {node.title for node in roots}
    assert "三、索赔及服务问答" not in root_titles
    assert "四、应急措施" not in root_titles
