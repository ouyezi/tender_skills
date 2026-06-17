from docx import Document

from doc_chunk.extract.table_grid import logical_rows_from_physical, parse_physical_grid


def test_parse_physical_grid_colspan(merged_colspan_docx):
    doc = Document(merged_colspan_docx)
    grid_width, rows = parse_physical_grid(doc.tables[0])
    assert grid_width == 3
    assert rows[0].cells[0].text == "姓名"
    assert rows[0].cells[0].colspan == 2
    assert rows[0].cells[1].text == "角色"
    assert rows[0].cells[1].colspan == 1


def test_logical_rows_dedupes_colspan(merged_colspan_docx):
    doc = Document(merged_colspan_docx)
    grid_width, rows = parse_physical_grid(doc.tables[0])
    logical = logical_rows_from_physical(rows)
    assert logical[0] == ["姓名", "角色"]
    assert logical[1] == ["刘敏", "开发"]
