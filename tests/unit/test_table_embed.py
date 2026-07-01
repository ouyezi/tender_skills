from __future__ import annotations

from pathlib import Path

from docx import Document

from doc_chunk.table.embed import embed_table_from_slice
from doc_chunk.table.slice_pack import build_mini_docx_for_table


def test_embed_table_from_slice_inserts_word_table(
    sample_docx_with_styled_table: Path, tmp_path: Path
) -> None:
    source = Document(sample_docx_with_styled_table)
    slice_path = tmp_path / "slice.docx"
    build_mini_docx_for_table(source.tables[0], slice_path)

    target = Document()
    anchor = target.add_paragraph("<!-- table-ref:tables/t0000.json -->")
    target.add_paragraph("| H | V |")

    embed_table_from_slice(target, slice_path, anchor)

    assert len(target.tables) == 1
    assert target.tables[0].cell(0, 0).text == "Header"
