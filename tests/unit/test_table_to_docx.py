from __future__ import annotations

from io import BytesIO

from docx import Document

from doc_chunk.convert.table_to_docx import render_sidecar_to_docx, render_table_to_docx
from doc_chunk.extract.table_extractor import extract_table
from doc_chunk.extract.table_grid import logical_rows_from_physical, parse_physical_grid
from doc_chunk.models.table_model import TableCell, TableGridRow


def _grid_dict(rows: list[TableGridRow]) -> dict:
    return {"rows": [r.model_dump() for r in rows]}


def test_render_table_roundtrip_merged_colspan(merged_colspan_docx) -> None:
    source = Document(merged_colspan_docx)
    grid_width, physical_rows = parse_physical_grid(source.tables[0])

    out_doc = Document()
    rendered = render_table_to_docx(out_doc, _grid_dict(physical_rows), grid_width=grid_width)

    out_width, out_rows = parse_physical_grid(rendered)
    assert out_width == grid_width
    assert logical_rows_from_physical(out_rows) == logical_rows_from_physical(physical_rows)


def test_render_table_roundtrip_personnel(personnel_dual_row_docx) -> None:
    source = Document(personnel_dual_row_docx)
    grid_width, physical_rows = parse_physical_grid(source.tables[0])

    out_doc = Document()
    rendered = render_table_to_docx(out_doc, _grid_dict(physical_rows), grid_width=grid_width)

    out_width, out_rows = parse_physical_grid(rendered)
    assert out_width == grid_width
    assert logical_rows_from_physical(out_rows) == logical_rows_from_physical(physical_rows)


def test_render_sidecar_updates_records(personnel_dual_row_docx) -> None:
    source = Document(personnel_dual_row_docx)
    markdown, sidecar, _ = extract_table(source.tables[0], block_index=0)
    assert sidecar is not None

    updated = [{**sidecar.records[0], "姓名": "张三", "级别": "架构师"}]
    out_doc = Document()
    rendered = render_sidecar_to_docx(out_doc, sidecar, records=updated)

    _, out_rows = parse_physical_grid(rendered)
    logical = logical_rows_from_physical(out_rows)
    assert logical[1][0] == "张三"
    assert logical[3][0] == "架构师"


def test_render_sidecar_roundtrip_via_bytes(personnel_dual_row_docx) -> None:
    source = Document(personnel_dual_row_docx)
    markdown, sidecar, _ = extract_table(source.tables[0], block_index=0)
    assert sidecar is not None

    out_doc = Document()
    render_sidecar_to_docx(out_doc, sidecar)
    buffer = BytesIO()
    out_doc.save(buffer)

    reloaded = Document(buffer)
    _, out_rows = parse_physical_grid(reloaded.tables[0])
    assert logical_rows_from_physical(out_rows) == sidecar.logical_rows


def test_render_table_empty_grid_raises() -> None:
    doc = Document()
    try:
        render_table_to_docx(doc, {"rows": []}, grid_width=1)
        raised = False
    except ValueError:
        raised = True
    assert raised


def test_render_simple_table_with_records() -> None:
    physical_rows = [
        TableGridRow(cells=[TableCell(text="序号"), TableCell(text="名称")]),
        TableGridRow(cells=[TableCell(text="1"), TableCell(text="foo")]),
    ]
    logical = [["序号", "名称"], ["1", "foo"]]
    doc = Document()
    table = render_table_to_docx(
        doc,
        _grid_dict(physical_rows),
        grid_width=2,
        records=[{"序号": "9", "名称": "bar"}],
        layout_type="simple",
        logical_rows=logical,
    )
    assert table.cell(1, 0).text == "9"
    assert table.cell(1, 1).text == "bar"
