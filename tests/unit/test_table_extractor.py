from docx import Document

from doc_chunk.extract.table_extractor import extract_table
from doc_chunk.extract.table_sidecar import TableSidecarWriter
from doc_chunk.models.table_model import TablesIndex
from doc_chunk.workspace.layout import OutputWorkspace


def test_extract_table_merged_colspan(merged_colspan_docx):
    doc = Document(merged_colspan_docx)
    markdown, sidecar, warnings = extract_table(doc.tables[0], block_index=0)

    assert warnings == []
    assert "| 姓名 | 角色 |" in markdown
    assert "| 刘敏 | 开发 |" in markdown
    assert "姓名 | 姓名" not in markdown
    assert sidecar is not None
    assert sidecar.logical_rows == [["姓名", "角色"], ["刘敏", "开发"]]
    assert sidecar.grid_width == 3
    assert sidecar.layout_type == "key_value"


def test_extract_table_personnel_dual_row(personnel_dual_row_docx):
    doc = Document(personnel_dual_row_docx)
    markdown, sidecar, warnings = extract_table(doc.tables[0], block_index=1)

    assert warnings == []
    assert sidecar is not None
    assert sidecar.layout_type == "personnel_dual_row"
    assert sidecar.record_groups == [[0, 1, 2, 3]]
    assert len(sidecar.records) == 1
    assert sidecar.records[0]["姓名"] == "刘敏"
    assert sidecar.records[0]["级别"] == "高级Java工程师"
    assert "【表格:人员信息】" in sidecar.llm_text
    assert "姓名: 刘敏" in sidecar.llm_text
    assert "| 姓名 |" in markdown


def test_table_sidecar_writer(tmp_path):
    ws = OutputWorkspace.create(tmp_path / "ws", overwrite=False)
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"

    _, sidecar, _ = extract_table(table, block_index=0)
    assert sidecar is not None

    writer = TableSidecarWriter(ws)
    rel = writer.write(sidecar)
    writer.finalize()

    assert rel == "tables/t0000.json"
    assert (ws.root / rel).is_file()
    index = TablesIndex.model_validate_json(ws.tables_index_path.read_text(encoding="utf-8"))
    assert len(index.tables) == 1
    assert index.tables[0].path == "tables/t0000.json"
