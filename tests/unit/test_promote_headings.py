from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from doc_chunk.api import extract_file, extract_outline
from doc_chunk.chunk.planner import plan_chunks
from doc_chunk.models.outline import OutlineTree

FIXTURES = Path(__file__).resolve().parents[1] / "fixtures"
NO_HEADING_DOCX = FIXTURES / "no_heading_style.docx"


@pytest.fixture(scope="module")
def no_heading_docx() -> Path:
    if NO_HEADING_DOCX.exists():
        return NO_HEADING_DOCX
    FIXTURES.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_paragraph("1. 技术方案")
    doc.add_paragraph("方案正文。")
    doc.add_paragraph("2. 报价说明")
    doc.add_paragraph("报价内容。")
    doc.save(NO_HEADING_DOCX)
    return NO_HEADING_DOCX


def test_promote_headings_auto_adds_hash_prefix(no_heading_docx: Path, tmp_path: Path) -> None:
    ws_plain = tmp_path / "plain"
    ws_promoted = tmp_path / "promoted"
    extract_file(no_heading_docx, ws_plain, overwrite=True, promote_headings="off")
    extract_file(no_heading_docx, ws_promoted, overwrite=True, promote_headings="auto")

    plain_md = (ws_plain / "content.md").read_text(encoding="utf-8")
    promoted_md = (ws_promoted / "content.md").read_text(encoding="utf-8")
    assert "# 技术方案" not in plain_md
    assert "# 技术方案" in promoted_md
    assert "# 报价说明" in promoted_md

    extract_outline(ws_plain)
    extract_outline(ws_promoted)
    plain_tree = OutlineTree.model_validate_json((ws_plain / "outline.json").read_text(encoding="utf-8"))
    promoted_tree = OutlineTree.model_validate_json((ws_promoted / "outline.json").read_text(encoding="utf-8"))
    plain_chunks = plan_chunks((ws_plain / "content.md").read_text(encoding="utf-8"), plain_tree)
    promoted_chunks = plan_chunks((ws_promoted / "content.md").read_text(encoding="utf-8"), promoted_tree)
    assert len([c for c in plain_chunks if c.title != "Preface"]) >= 2
    assert len(plain_chunks) == len(promoted_chunks)
