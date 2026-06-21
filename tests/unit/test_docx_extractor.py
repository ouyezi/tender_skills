from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from doc_chunk.extract.docx_extractor import extract_docx
from doc_chunk.models.content_block import ContentBlocksFile
from doc_chunk.models.table_model import TableSidecar, TablesIndex
from doc_chunk.workspace.layout import OutputWorkspace


def _set_outline_level(paragraph, level: int) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    existing = p_pr.find(qn("w:outlineLvl"))
    if existing is not None:
        p_pr.remove(existing)
    outline_lvl = OxmlElement("w:outlineLvl")
    outline_lvl.set(qn("w:val"), str(level - 1))
    p_pr.append(outline_lvl)


def test_extract_docx_writes_markdown(sample_docx: Path, tmp_path: Path) -> None:
    ws = OutputWorkspace.create(tmp_path / "ws", overwrite=False)

    result = extract_docx(sample_docx, ws)

    content = ws.content_path.read_text(encoding="utf-8")
    assert ws.content_path.exists()
    assert "# 示例标题" in content
    assert "这是一段用于测试的正文。" in content
    assert result.image_count == 0
    assert result.warnings == []


def test_extract_docx_deduplicates_repeated_embed_references(
    sample_docx_with_duplicate_embed: Path, tmp_path: Path
) -> None:
    ws = OutputWorkspace.create(tmp_path / "ws", overwrite=False)

    result = extract_docx(sample_docx_with_duplicate_embed, ws)

    content = ws.content_path.read_text(encoding="utf-8")
    images = [p for p in ws.images_dir.iterdir() if p.name != "manifest.json"]
    assert result.image_count == 2
    assert len(images) == 1
    assert content.count("![") == 2
    assert content.count("images/docx-img-001") == 2


def test_extract_docx_deduplicates_repeated_embed_in_paragraph(
    sample_docx_with_duplicate_blip_in_paragraph: Path, tmp_path: Path
) -> None:
    ws = OutputWorkspace.create(tmp_path / "ws", overwrite=False)

    result = extract_docx(sample_docx_with_duplicate_blip_in_paragraph, ws)

    content = ws.content_path.read_text(encoding="utf-8")
    images = [p for p in ws.images_dir.iterdir() if p.name != "manifest.json"]
    assert result.image_count == 1
    assert len(images) == 1
    assert content.count("![") == 1


def test_extract_docx_reuses_file_across_paragraphs(
    sample_docx_with_reused_image_across_paragraphs: Path, tmp_path: Path
) -> None:
    ws = OutputWorkspace.create(tmp_path / "ws", overwrite=False)

    result = extract_docx(sample_docx_with_reused_image_across_paragraphs, ws)

    content = ws.content_path.read_text(encoding="utf-8")
    images = [p for p in ws.images_dir.iterdir() if p.name != "manifest.json"]
    assert result.image_count == 2
    assert len(images) == 1
    assert content.count("![") == 2
    assert "![docx-img-001]" in content
    assert "![docx-img-002]" in content
    assert content.count("images/docx-img-001") == 2


def test_extract_docx_table_reuses_rids_from_body(
    sample_docx_with_reused_image_in_body_and_table: Path, tmp_path: Path
) -> None:
    ws = OutputWorkspace.create(tmp_path / "ws", overwrite=False)

    result = extract_docx(sample_docx_with_reused_image_in_body_and_table, ws)

    content = ws.content_path.read_text(encoding="utf-8")
    images = [p for p in ws.images_dir.iterdir() if p.name != "manifest.json"]
    assert result.image_count == 2
    assert len(images) == 1
    assert content.count("![") == 2
    assert content.count("images/docx-img-001") == 2


def test_extract_docx_exports_inline_images(
    sample_docx_with_image: Path, tmp_path: Path
) -> None:
    ws = OutputWorkspace.create(tmp_path / "ws", overwrite=False)

    result = extract_docx(sample_docx_with_image, ws)

    content = ws.content_path.read_text(encoding="utf-8")
    images = [p for p in ws.images_dir.iterdir() if p.name != "manifest.json"]
    assert result.image_count == 1
    assert len(images) == 1
    assert images[0].name.startswith("docx-img-001")
    assert "![docx-img-001]" in content


def test_extract_docx_recognizes_outline_level_without_heading_style(tmp_path: Path) -> None:
    docx_path = tmp_path / "outline-level.docx"
    doc = Document()
    doc.add_paragraph("12.信用中国无列入失信被执行人证明", style="Heading 2")
    body_heading = doc.add_paragraph("13.信用中国无重大税收违法失信主体证明")
    _set_outline_level(body_heading, 2)
    doc.add_paragraph("证明图片说明")
    para_14 = doc.add_paragraph("14. 政府采购无严重违法失信行为证明")
    _set_outline_level(para_14, 2)
    doc.add_paragraph("15.荣誉证书", style="Heading 2")
    doc.save(docx_path)

    ws = OutputWorkspace.create(tmp_path / "ws", overwrite=False)
    extract_docx(docx_path, ws)

    content = ws.content_path.read_text(encoding="utf-8")
    assert "## 12.信用中国无列入失信被执行人证明" in content
    assert "## 13.信用中国无重大税收违法失信主体证明" in content
    assert "## 14. 政府采购无严重违法失信行为证明" in content
    assert "## 15.荣誉证书" in content
    assert content.index("## 13.") < content.index("## 14.")
    assert content.index("## 14.") < content.index("## 15.")


def test_extract_docx_table_sidecar(merged_colspan_docx: Path, tmp_path: Path) -> None:
    ws = OutputWorkspace.create(tmp_path / "ws", overwrite=False)

    extract_docx(merged_colspan_docx, ws)

    assert ws.tables_index_path.exists()
    index = TablesIndex.model_validate_json(ws.tables_index_path.read_text(encoding="utf-8"))
    assert len(index.tables) == 1
    assert index.tables[0].path == "tables/t0000.json"

    blocks = ContentBlocksFile.model_validate_json(ws.content_blocks_path.read_text(encoding="utf-8"))
    assert blocks.schema_version == "1.1"
    table_blocks = [b for b in blocks.blocks if b.block_type == "table"]
    assert len(table_blocks) == 1
    assert table_blocks[0].table_ref == "tables/t0000.json"

    sidecar = TableSidecar.model_validate_json((ws.tables_dir / "t0000.json").read_text(encoding="utf-8"))
    assert sidecar.block_index == 0
    assert sidecar.markdown

    content = ws.content_path.read_text(encoding="utf-8")
    assert "姓名 | 姓名" not in content
    assert "刘敏 | 刘敏" not in content
