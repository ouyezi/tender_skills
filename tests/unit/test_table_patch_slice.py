from __future__ import annotations

from pathlib import Path

from docx import Document

from doc_chunk.extract.docx_extractor import extract_docx
from doc_chunk.table.patch import patch_docx_tables
from doc_chunk.workspace.layout import OutputWorkspace


def test_patch_docx_tables_uses_slice_when_ok(
    sample_docx_with_styled_table: Path, tmp_path: Path
) -> None:
    ws = OutputWorkspace.create(tmp_path / "ws", overwrite=True)
    extract_docx(sample_docx_with_styled_table, ws)

    md = ws.content_path.read_text(encoding="utf-8")
    doc = Document()
    for line in md.splitlines():
        doc.add_paragraph(line)

    result = patch_docx_tables(doc, ws)
    assert result.patched_count == 1
    assert len(doc.tables) == 1
    assert doc.tables[0].cell(0, 0).text == "Header"
