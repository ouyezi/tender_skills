from docx import Document

from doc_chunk.extract.table_grid import logical_rows_from_physical, parse_physical_grid
from doc_chunk.extract.table_layout import build_records, classify_layout


def test_classify_personnel_dual_row(personnel_dual_row_docx):
    doc = Document(personnel_dual_row_docx)
    _, rows = parse_physical_grid(doc.tables[0])
    logical = logical_rows_from_physical(rows)
    layout, groups = classify_layout(logical)
    assert layout == "personnel_dual_row"
    assert groups == [[0, 1, 2, 3]]


def test_build_records_merges_personnel_block(personnel_dual_row_docx):
    doc = Document(personnel_dual_row_docx)
    _, rows = parse_physical_grid(doc.tables[0])
    logical = logical_rows_from_physical(rows)
    layout, groups = classify_layout(logical)
    records = build_records(logical, layout, groups)
    assert len(records) == 1
    assert records[0]["姓名"] == "刘敏"
    assert records[0]["级别"] == "高级Java工程师"
    assert records[0]["毕业学校"] == "承德石油学院"


def test_classify_simple():
    logical = [["序号", "名称"], ["1", "foo"], ["2", "bar"]]
    layout, _ = classify_layout(logical)
    assert layout == "simple"
    records = build_records(logical, layout, [])
    assert records[0]["序号"] == "1"
