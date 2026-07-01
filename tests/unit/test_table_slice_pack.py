from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document

from doc_chunk.table.slice_pack import build_mini_docx_for_table

_STYLE_PARTS = (
    "word/styles.xml",
    "word/stylesWithEffects.xml",
    "word/theme/theme1.xml",
    "word/fontTable.xml",
    "word/numbering.xml",
    "word/settings.xml",
)


def test_build_mini_docx_for_table_writes_valid_docx(
    sample_docx_with_styled_table: Path, tmp_path: Path
) -> None:
    source = Document(sample_docx_with_styled_table)
    dest = tmp_path / "slice.docx"
    build_mini_docx_for_table(source.tables[0], dest)

    assert dest.is_file()
    assert dest.read_bytes()[:2] == b"PK"

    loaded = Document(dest)
    assert len(loaded.tables) == 1
    assert loaded.tables[0].cell(0, 0).text == "Header"
    assert loaded.tables[0].cell(1, 1).text == "B"


def test_build_mini_docx_preserves_colspan(merged_colspan_docx: Path, tmp_path: Path) -> None:
    source = Document(merged_colspan_docx)
    dest = tmp_path / "merged_slice.docx"
    build_mini_docx_for_table(source.tables[0], dest)

    loaded = Document(dest)
    assert len(loaded.tables) == 1
    assert "姓名" in loaded.tables[0].cell(0, 0).text


def test_build_mini_docx_preserves_cell_image(
    sample_docx_with_reused_image_in_body_and_table: Path, tmp_path: Path
) -> None:
    source = Document(sample_docx_with_reused_image_in_body_and_table)
    dest = tmp_path / "img_slice.docx"
    build_mini_docx_for_table(source.tables[0], dest)

    loaded = Document(dest)
    assert len(loaded.tables) == 1
    blips = [node for node in loaded.tables[0]._tbl.iter() if node.tag.endswith("}blip")]
    assert len(blips) >= 1


def test_build_mini_docx_grafts_style_parts_from_source(
    sample_docx_with_styled_table: Path, tmp_path: Path
) -> None:
    source = Document(sample_docx_with_styled_table)
    slice_path = tmp_path / "slice.docx"
    build_mini_docx_for_table(source.tables[0], slice_path)

    with zipfile.ZipFile(sample_docx_with_styled_table) as src_zip, zipfile.ZipFile(slice_path) as slice_zip:
        for part_path in _STYLE_PARTS:
            if part_path in src_zip.namelist() and part_path in slice_zip.namelist():
                assert slice_zip.read(part_path) == src_zip.read(part_path), part_path
