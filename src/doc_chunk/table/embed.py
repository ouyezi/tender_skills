from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

from doc_chunk.table.slice_graft import graft_style_blobs, remap_embed_relationships


def embed_table_from_slice(
    target_doc: Document,
    slice_path: Path,
    before_paragraph: Paragraph,
) -> None:
    slice_doc = Document(slice_path)
    if not slice_doc.tables:
        raise ValueError(f"slice has no table: {slice_path}")

    tbl_copy = deepcopy(slice_doc.tables[0]._tbl)
    graft_style_blobs(slice_doc, target_doc)
    remap_embed_relationships(slice_doc.part, target_doc.part, tbl_copy)
    before_paragraph._p.addprevious(tbl_copy)
