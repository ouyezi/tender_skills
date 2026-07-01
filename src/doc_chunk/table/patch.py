from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

from docx.document import Document
from docx.text.paragraph import Paragraph

from doc_chunk.convert.table_to_docx import render_sidecar_to_docx
from doc_chunk.models.table_model import TableSidecar
from doc_chunk.table.access import load_table_model
from doc_chunk.table.embed import embed_table_from_slice
from doc_chunk.table.placeholders import parse_table_ref_from_line
from doc_chunk.workspace.layout import OutputWorkspace


@dataclass
class PatchResult:
    patched_count: int = 0
    skipped: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


def _is_markdown_table_line(text: str) -> bool:
    stripped = text.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def _collect_markdown_table_paragraphs(paragraphs: list[Paragraph], start_index: int) -> list[int]:
    indices: list[int] = []
    for idx in range(start_index + 1, len(paragraphs)):
        text = paragraphs[idx].text
        if _is_markdown_table_line(text):
            indices.append(idx)
            continue
        if not text.strip():
            indices.append(idx)
            continue
        break
    return indices


def _insert_table_from_sidecar(
    document: Document,
    paragraph: Paragraph,
    sidecar: TableSidecar,
    ws: OutputWorkspace,
    result: PatchResult,
) -> None:
    if sidecar.slice_status == "ok" and sidecar.slice_ref:
        slice_path = ws.root / sidecar.slice_ref
        if slice_path.is_file():
            try:
                embed_table_from_slice(document, slice_path, paragraph)
                return
            except Exception as exc:
                result.warnings.append(f"table_embed_fallback:{sidecar.slice_ref}:{exc}")
    table = render_sidecar_to_docx(document, sidecar)
    paragraph._p.addprevious(table._tbl)


def patch_docx_tables(
    document: Document,
    workspace: OutputWorkspace | Path,
    *,
    table_refs: list[str] | None = None,
) -> PatchResult:
    ws = workspace if isinstance(workspace, OutputWorkspace) else OutputWorkspace.open_existing(Path(workspace))
    allowed = set(table_refs) if table_refs else None
    paragraphs = list(document.paragraphs)
    result = PatchResult()
    targets: list[tuple[int, str, list[int]]] = []

    for idx, paragraph in enumerate(paragraphs):
        ref = parse_table_ref_from_line(paragraph.text)
        if ref is None:
            continue
        if allowed is not None and ref not in allowed:
            continue
        md_indices = _collect_markdown_table_paragraphs(paragraphs, idx)
        targets.append((idx, ref, md_indices))

    for placeholder_idx, ref, md_indices in reversed(targets):
        sidecar_path = ws.root / ref
        if not sidecar_path.is_file():
            result.skipped.append(ref)
            result.warnings.append(f"table_sidecar_missing:{ref}")
            continue
        sidecar = load_table_model(ws, ref)
        placeholder = document.paragraphs[placeholder_idx]
        _insert_table_from_sidecar(document, placeholder, sidecar, ws, result)
        delete_indices = sorted([placeholder_idx] + md_indices, reverse=True)
        for del_idx in delete_indices:
            p = document.paragraphs[del_idx]
            p._element.getparent().remove(p._element)
        result.patched_count += 1

    return result
