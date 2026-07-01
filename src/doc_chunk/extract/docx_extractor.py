from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph

from typing import Literal

from doc_chunk.extract.block_index import BlockAccumulator, write_accumulator_markdown, write_content_blocks
from doc_chunk.extract.docx_numbering import DocxNumberingResolver, merge_list_prefix
from doc_chunk.extract.promote_headings import parse_content_heading_line
from doc_chunk.extract.table_extractor import extract_table
from doc_chunk.extract.table_sidecar import TableSidecarWriter
from doc_chunk.extract.table_slice import extract_table_slice
from doc_chunk.models.document import ExtractResult
from doc_chunk.models.images_manifest import ImageManifestEntry, ImagesManifest
from doc_chunk.table.assets import collect_table_assets
from doc_chunk.workspace.layout import OutputWorkspace


def _heading_level_from_style(style_name: str) -> int | None:
    name = style_name.strip()
    parts = name.split()
    if len(parts) == 2 and parts[1].isdigit():
        if parts[0] in {"Heading", "标题"}:
            return max(1, min(8, int(parts[1])))
    return None


def _outline_level_from_paragraph(paragraph: DocxParagraph) -> int | None:
    p_pr = paragraph._element.pPr
    if p_pr is None:
        return None
    outline_lvl = p_pr.find(qn("w:outlineLvl"))
    if outline_lvl is None:
        return None
    raw = outline_lvl.get(qn("w:val"))
    if raw is None:
        return None
    try:
        level = int(raw) + 1
    except ValueError:
        return None
    if not 1 <= level <= 8:
        return None
    return level


def _resolve_paragraph_heading_level(paragraph: DocxParagraph, text: str) -> int | None:
    level = _heading_level_from_style(paragraph.style.name)
    if level is not None:
        return level

    outline_level = _outline_level_from_paragraph(paragraph)
    if outline_level is not None and parse_content_heading_line(text) is not None:
        return outline_level

    return None


def _docx_element_image_embeds(element: object, doc: DocxDocument) -> list[tuple[str, object]]:
    embeds: list[tuple[str, object]] = []
    seen_relationship_ids: set[str] = set()
    for child in element.iter():
        if child.tag != qn("a:blip"):
            continue
        relationship_id = child.get(qn("r:embed"))
        if relationship_id is None or relationship_id in seen_relationship_ids:
            continue
        image_part = doc.part.related_parts.get(relationship_id)
        if image_part is not None:
            seen_relationship_ids.add(relationship_id)
            embeds.append((relationship_id, image_part))
    return embeds


def _docx_paragraph_image_embeds(paragraph: DocxParagraph, doc: DocxDocument) -> list[tuple[str, object]]:
    return _docx_element_image_embeds(paragraph._element, doc)


def _docx_table_image_embeds(table: DocxTable, doc: DocxDocument) -> list[tuple[str, object]]:
    embeds: list[tuple[str, object]] = []
    seen_cells: set[int] = set()
    for row in table.rows:
        for cell in row.cells:
            cell_id = id(cell._tc)
            if cell_id in seen_cells:
                continue
            seen_cells.add(cell_id)
            for paragraph in cell.paragraphs:
                embeds.extend(_docx_paragraph_image_embeds(paragraph, doc))
    return embeds


def _image_extension(partname: object, content_type: str) -> str:
    suffix = Path(str(partname)).suffix.lower()
    if suffix:
        return suffix
    return {
        "image/png": ".png",
        "image/jpeg": ".jpg",
        "image/webp": ".webp",
        "image/gif": ".gif",
    }.get(content_type.lower(), ".bin")


def _save_docx_image(workspace: OutputWorkspace, image_part: object, image_number: int) -> str:
    extension = _image_extension(image_part.partname, image_part.content_type)
    image_name = f"docx-img-{image_number:03d}{extension}"
    (workspace.images_dir / image_name).write_bytes(image_part.blob)
    return image_name


def _register_docx_image(
    *,
    relationship_id: str,
    image_part: object,
    relationship_to_image_ref: dict[str, str],
    workspace: OutputWorkspace,
    acc: BlockAccumulator,
    image_entries: list[ImageManifestEntry],
    image_count: int,
) -> int:
    if relationship_id not in relationship_to_image_ref:
        file_number = len(relationship_to_image_ref) + 1
        image_name = _save_docx_image(workspace, image_part, file_number)
        relationship_to_image_ref[relationship_id] = f"images/{image_name}"
    else:
        image_name = Path(relationship_to_image_ref[relationship_id]).name

    image_count += 1
    image_ref = relationship_to_image_ref[relationship_id]
    block_index_before = acc.block_count
    acc.add_image(image_ref, alt=f"docx-img-{image_count:03d}")
    image_entries.append(
        ImageManifestEntry(
            image_ref=image_ref,
            file_name=image_name,
            content_type=image_part.content_type,
            byte_size=len(image_part.blob),
            source_block_index=block_index_before,
        )
    )
    return image_count


def extract_docx(
    path: Path,
    workspace: OutputWorkspace,
    *,
    promote_headings: Literal["off", "auto"] = "off",
) -> ExtractResult:
    doc = DocxDocument(path)
    numbering = DocxNumberingResolver(doc)
    acc = BlockAccumulator()
    sidecar_writer = TableSidecarWriter(workspace)
    image_count = 0
    image_entries: list[ImageManifestEntry] = []
    relationship_to_image_ref: dict[str, str] = {}
    all_warnings: list[str] = []

    for element in doc.element.body.iterchildren():
        if element.tag.endswith("}p"):
            paragraph = DocxParagraph(element, doc)
            list_prefix = numbering.advance(paragraph)
            text = paragraph.text.strip()
            if text and list_prefix:
                text = merge_list_prefix(text, list_prefix)
            if text:
                level = _resolve_paragraph_heading_level(paragraph, text)
                if level is not None:
                    acc.add_heading(level, text)
                elif promote_headings == "auto":
                    parsed = parse_content_heading_line(text)
                    if parsed is not None:
                        acc.add_heading(parsed[0], parsed[1])
                    else:
                        acc.add_paragraph(text)
                else:
                    acc.add_paragraph(text)
            for relationship_id, image_part in _docx_paragraph_image_embeds(paragraph, doc):
                image_count = _register_docx_image(
                    relationship_id=relationship_id,
                    image_part=image_part,
                    relationship_to_image_ref=relationship_to_image_ref,
                    workspace=workspace,
                    acc=acc,
                    image_entries=image_entries,
                    image_count=image_count,
                )
            continue

        if element.tag.endswith("}tbl"):
            table = DocxTable(element, doc)
            block_index_before = acc.block_count
            markdown, sidecar, tbl_warnings = extract_table(table, block_index_before)
            all_warnings.extend(tbl_warnings)
            table_ref = None
            if markdown:
                if sidecar:
                    slice_ref, slice_status, slice_warnings = extract_table_slice(
                        table, sidecar.block_index, workspace.root
                    )
                    all_warnings.extend(slice_warnings)
                    table_ref = sidecar_writer.write(
                        sidecar,
                        slice_ref=slice_ref,
                        slice_status=slice_status,
                    )
                acc.add_table(markdown, table_ref=table_ref)
            for relationship_id, image_part in _docx_table_image_embeds(table, doc):
                image_count = _register_docx_image(
                    relationship_id=relationship_id,
                    image_part=image_part,
                    relationship_to_image_ref=relationship_to_image_ref,
                    workspace=workspace,
                    acc=acc,
                    image_entries=image_entries,
                    image_count=image_count,
                )

    write_accumulator_markdown(workspace, acc)
    write_content_blocks(workspace, acc.finalize())
    sidecar_writer.finalize()
    if image_entries:
        manifest = ImagesManifest(images=image_entries)
        workspace.images_manifest_path.write_text(manifest.model_dump_json(indent=2), encoding="utf-8")
    collect_table_assets(workspace, write_manifest=True)
    return ExtractResult(image_count=image_count, warnings=all_warnings)
