from __future__ import annotations

import tempfile
from pathlib import Path

from doc_chunk.api import run_pipeline
from doc_chunk.models.outline import Anchor, OutlineNode, OutlineTree
from docx import Document

from viewer.services.outline_tree import PREFACE_NODE_ID
from viewer.services.section_slice import slice_section


def test_slice_section_uses_markdown_headings_not_block_anchors() -> None:
    content_md = "Preface\n\n# Chapter 1\n\nAlpha\n\n## Section 1.1\n\nBeta\n\n# Chapter 2\n\nGamma"
    tree = OutlineTree(
        nodes=[
            OutlineNode(
                node_id="n1",
                title="Chapter 1",
                level=1,
                parent_id=None,
                sort_order=0,
                anchor=Anchor(char_start=content_md.index("Alpha")),
            ),
            OutlineNode(
                node_id="n2",
                title="Section 1.1",
                level=2,
                parent_id="n1",
                sort_order=1,
                anchor=Anchor(char_start=content_md.index("Alpha")),
            ),
        ]
    )

    preface = slice_section(content_md, tree, PREFACE_NODE_ID)
    assert preface.title == "前言"
    assert "Preface" in preface.markdown
    assert "# Chapter 1" not in preface.markdown

    chapter = slice_section(content_md, tree, "n1")
    assert "Alpha" in chapter.markdown
    assert "Beta" in chapter.markdown
    assert "Gamma" not in chapter.markdown
    assert chapter.markdown.startswith("# Chapter 1")

    section = slice_section(content_md, tree, "n2")
    assert section.title == "Section 1.1"
    assert "Beta" in section.markdown
    assert "Gamma" not in section.markdown
    assert "Alpha" not in section.markdown
    assert section.section_path == ["Chapter 1", "Section 1.1"]


def test_slice_section_aligns_with_pipeline_outline() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        docx = root / "sample.docx"
        doc = Document()
        doc.add_heading("第一章", level=1)
        doc.add_paragraph("第一章引言段落")
        doc.add_heading("1.1 小节", level=2)
        doc.add_paragraph("小节正文A")
        doc.add_heading("1.2 小节", level=2)
        doc.add_paragraph("小节正文B")
        doc.add_heading("第二章", level=1)
        doc.add_paragraph("第二章正文")
        doc.save(docx)

        workspace = root / "ws"
        result = run_pipeline(docx, workspace, overwrite=True, skip_refine=True, skip_enrich=True)
        assert result.status == "success"

        content_md = (workspace / "content.md").read_text(encoding="utf-8")
        outline = OutlineTree.model_validate_json((workspace / "outline.json").read_text(encoding="utf-8"))

        chapter_one = next(node for node in outline.nodes if node.title == "第一章")
        section_one_one = next(node for node in outline.nodes if node.title == "1.1 小节")
        chapter_two = next(node for node in outline.nodes if node.title == "第二章")

        ch1 = slice_section(content_md, outline, chapter_one.node_id)
        assert "第一章引言段落" in ch1.markdown
        assert "小节正文A" in ch1.markdown
        assert "小节正文B" in ch1.markdown
        assert "第二章正文" not in ch1.markdown

        s11 = slice_section(content_md, outline, section_one_one.node_id)
        assert s11.markdown.startswith("## 1.1 小节")
        assert "小节正文A" in s11.markdown
        assert "1.2 小节" not in s11.markdown.split("小节正文A", 1)[-1]

        ch2 = slice_section(content_md, outline, chapter_two.node_id)
        assert "第二章正文" in ch2.markdown
        assert "第一章" not in ch2.markdown
